"""DOCX file parser for ATG."""
from pathlib import Path
from typing import List, Optional

import docx

from . import BaseParser


class DocxParser(BaseParser):
    """Parser for Microsoft Word documents (.docx)."""

    @classmethod
    def supported_formats(cls) -> tuple[str, ...]:
        """Get the file extensions supported by this parser.

        Returns:
            A tuple of supported file extensions.
        """
        return (".docx",)

    def _get_paragraphs_text(self, paragraphs: List) -> List[str]:
        """Extract text from a list of paragraphs.

        Args:
            paragraphs: List of paragraph objects from python-docx.

        Returns:
            List of paragraph texts.
        """
        return [p.text for p in paragraphs if p.text.strip()]

    def _get_tables_text(self, tables: List) -> List[str]:
        """Extract text from tables.

        Args:
            tables: List of table objects from python-docx.

        Returns:
            List of table texts.
        """
        table_texts = []
        for table in tables:
            rows_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                rows_text.append(" | ".join(row_text))
            if rows_text:
                table_texts.append("\n".join(rows_text))
        return table_texts

    def _get_header_footer_text(self, doc) -> List[str]:
        """Extract text from headers and footers.

        Args:
            doc: The document object from python-docx.

        Returns:
            List of header and footer texts.
        """
        texts = []
        for section in doc.sections:
            # Get header text
            if section.header is not None:
                texts.extend(self._get_paragraphs_text(section.header.paragraphs))
            # Get footer text
            if section.footer is not None:
                texts.extend(self._get_paragraphs_text(section.footer.paragraphs))
        return texts

    def parse(self, file_path: Path) -> str:
        """Parse a DOCX file and return its content as plain text.

        Args:
            file_path: Path to the DOCX file to parse.

        Returns:
            The parsed content as plain text.

        Raises:
            docx.opc.exceptions.PackageNotFoundError: If the file is not a valid DOCX.
            PermissionError: If the file cannot be accessed.
            ValueError: If there's an error parsing the document.
        """
        try:
            # Load the document
            doc = docx.Document(file_path)

            # Extract content from different parts of the document
            paragraphs = self._get_paragraphs_text(doc.paragraphs)
            tables = self._get_tables_text(doc.tables)
            headers_footers = self._get_header_footer_text(doc)

            # Combine all content
            content_parts = []
            if paragraphs:
                content_parts.append("\n\n".join(paragraphs))
            if tables:
                content_parts.append("\n\nTables:\n" + "\n\n".join(tables))
            if headers_footers:
                content_parts.append(
                    "\n\nHeaders/Footers:\n" + "\n".join(headers_footers)
                )

            return "\n\n".join(content_parts)

        except Exception as e:
            raise ValueError(f"Error parsing DOCX file {file_path}: {str(e)}")
